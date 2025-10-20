
import os, json, datetime, io
import azure.functions as func
from azure.search.documents import SearchClient
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from openai import AzureOpenAI
from azure.storage.blob import BlobClient
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SEARCH_ENDPOINT = os.environ["SEARCH_ENDPOINT"]
SEARCH_INDEX = os.environ.get("SEARCH_INDEX", "support-docs")
AOAI_ENDPOINT = os.environ["AOAI_ENDPOINT"]
AOAI_DEPLOYMENT = os.environ.get("AOAI_DEPLOYMENT", "gpt-4o-mini")
ST_CONN = os.environ["DATA_STORAGE_CONNECTION"]
RUNBOOKS_CONTAINER = os.getenv("RUNBOOKS_CONTAINER", "runbooks")

search = SearchClient(SEARCH_ENDPOINT, SEARCH_INDEX, DefaultAzureCredential())
cred = DefaultAzureCredential()
provider = get_bearer_token_provider(cred, "https://cognitiveservices.azure.com/.default")
aoai = AzureOpenAI(azure_endpoint=AOAI_ENDPOINT, api_version="2024-06-01", azure_ad_token_provider=provider)

SYSTEM_PROMPT = (
    "You are an expert Azure support runbook writer. Produce a single comprehensive runbook suitable for all levels "
    "in managed services without L1/L2/L3 labels. Use British English; be concise, operational, accurate. "
    "Use ONLY provided sources and add a 'Gaps & Assumptions' section when information is missing. "
    "Return a STRICT JSON structure with the following keys: "
    "title, metadata (customerId, serviceArea, generated, version), sections (array of {heading, body, lists?, tables?}), "
    "and sources (array of {label, path}). The 'body' should be plain paragraphs (no Markdown)."
)

app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)


def _build_docx(model_json: dict) -> Document:
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Segoe UI'
    styles['Normal'].font.size = Pt(10.5)

    title = model_json.get('title') or 'Support Runbook'
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = model_json.get('metadata', {})
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid'
    r0 = table.rows[0].cells
    r0[0].text = 'Customer'
    r0[1].text = meta.get('customerId','')
    r0[2].text = 'Service Area'
    r0[3].text = meta.get('serviceArea','')
    r1 = table.rows[1].cells
    r1[0].text = 'Generated'
    r1[1].text = meta.get('generated','')
    r1[2].text = 'Version'
    r1[3].text = meta.get('version','0.1')

    doc.add_paragraph('')

    for sec in model_json.get('sections', []):
        h = sec.get('heading')
        if h:
            doc.add_heading(h, level=1)
        body = sec.get('body') or ''
        if body:
            for para in body.split('\n'):
             para = para.strip()
    if para:
        doc.add_paragraph(para)
        for lst in sec.get('lists') or []:
            for item in lst:
                doc.add_paragraph(item, style='List Bullet')
        for t in sec.get('tables') or []:
            headers = t.get('headers') or []
            rows = t.get('rows') or []
            if headers:
                tbl = doc.add_table(rows=1, cols=len(headers))
                tbl.style = 'Light Grid'
                hdr_cells = tbl.rows[0].cells
                for i, hcell in enumerate(headers):
                    hdr_cells[i].text = str(hcell)
                for r in rows:
                    cells = tbl.add_row().cells
                    for i, v in enumerate(r):
                        if i < len(cells):
                            cells[i].text = str(v)
                doc.add_paragraph('')

    sources = model_json.get('sources') or []
    if sources:
        doc.add_heading('Appendix: Source Index', level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Light Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Label'
        hdr[1].text = 'Path'
        for s in sources:
            row = tbl.add_row().cells
            row[0].text = s.get('label','')
            row[1].text = s.get('path','')

    return doc


@app.function_name(name="generate")
@app.route(route="generate", methods=["POST"])
def run(req: func.HttpRequest) -> func.HttpResponse:
    body = req.get_json()
    cust = body.get("customerId")
    svc = body.get("serviceArea")
    query = body.get("query", f"{svc} runbook")
    if not cust or not svc:
        return func.HttpResponse("customerId and serviceArea required", status_code=400)

    results = search.search(search_text=query, filter=f"customerId eq '{cust}' and serviceArea eq '{svc}'", top=14)
    chunks = [{"text": r["text"], "source": r.get("sourceUrl") } for r in results]

    srcs = [{"label": f"S{i}", "path": c.get("source")} for i,c in enumerate(chunks, start=1)]

    user_payload = {
        "customerId": cust,
        "serviceArea": svc,
        "now": datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'),
        "sources": chunks
    }

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": json.dumps(user_payload)}
    ]

    resp = aoai.chat.completions.create(model=AOAI_DEPLOYMENT, messages=messages, temperature=0.2)
    raw = resp.choices[0].message.content

    try:
        model_json = json.loads(raw)
    except Exception:
        model_json = {
            "title": f"{cust} – {svc} Support Runbook",
            "metadata": {"customerId": cust, "serviceArea": svc, "generated": datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'), "version": "0.1"},
            "sections": [ {"heading": "Runbook", "body": raw} ],
            "sources": srcs
        }

    meta = model_json.setdefault('metadata', {})
    meta.setdefault('customerId', cust)
    meta.setdefault('serviceArea', svc)
    meta.setdefault('generated', datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'))
    meta.setdefault('version', '0.1')

    doc = _build_docx(model_json)

    stamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    path = f"{cust}/{svc}/{stamp}/runbook.docx"
    blob = BlobClient.from_connection_string(ST_CONN, container_name=RUNBOOKS_CONTAINER, blob_name=path)
    buf = io.BytesIO()
    doc.save(buf)
    blob.upload_blob(buf.getvalue(), overwrite=True)

    return func.HttpResponse(json.dumps({"status": "ok", "docxPath": f"{RUNBOOKS_CONTAINER}/{path}"}), mimetype="application/json")
